import io
import re
from dataclasses import dataclass
from typing import List, Tuple, Dict, Optional, Any

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches


# ============================
# Configuration
# ============================
APP_TITLE = "AutoChi: Automated Test for Proportations"
DEFAULT_PREVIEW_ROWS = 15

st.text("AutoChi is a Streamlit-based statistical analysis app for automated chi-square analysis of proportions with one or two categorical factors and two response-count columns.")
st.subheader("Expected input structure:")
st.markdown("""
The app assumes the uploaded CSV/XLSX file has this layout:
- First **1 or 2 columns**: categorical factor column(s)
- Last **2 columns**: response-count columns
- First row: variable names""")



@dataclass
class DetectionConfig:
    numeric_ratio_threshold: float = 0.85
    low_unique_ratio_threshold: float = 0.30
    min_non_null_ratio: float = 0.50


# ============================
# File + Detection Utilities
# ============================
def read_uploaded_file(uploaded_file) -> pd.DataFrame:
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        try:
            return pd.read_csv(uploaded_file)
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding="latin-1")
    if name.endswith(".xlsx"):
        return pd.read_excel(uploaded_file, engine="openpyxl")
    raise ValueError("Unsupported file type.")


def coerce_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def column_profile(series: pd.Series) -> Dict[str, float]:
    non_null = series.dropna()
    if len(series) == 0 or len(non_null) == 0:
        return {"non_null_ratio": 0, "numeric_ratio": 0, "unique_ratio": 0}

    return {
        "non_null_ratio": len(non_null) / len(series),
        "numeric_ratio": coerce_numeric(non_null).notna().mean(),
        "unique_ratio": non_null.astype(str).nunique() / len(non_null),
    }


def auto_detect_columns(df: pd.DataFrame, cfg: DetectionConfig):
    """
    Fixed layout assumption requested by the user:
    - the LAST TWO columns are always the two response-count columns
    - the one or two columns before them are the categorical factor column(s)

    The profiling dataframe is still returned for transparency/debugging,
    but it is no longer used to decide the defaults.
    """
    profs = []
    for c in df.columns:
        p = column_profile(df[c])
        profs.append({"column": c, **p})
    prof_df = pd.DataFrame(profs)

    if df.shape[1] < 3:
        raise ValueError(
            "Input file must have at least 3 columns: 1 factor + 2 response columns."
        )

    responses = df.columns[-2:].tolist()
    factors = df.columns[:-2].tolist()

    if len(factors) not in (1, 2):
        raise ValueError(
            "Input file must contain exactly 1 or 2 factor columns before the final 2 response columns."
        )

    return factors, responses, prof_df


# ============================
# Statistics Helpers
# ============================
def _normal_cdf(x):
    from math import erf, sqrt
    return 0.5 * (1 + erf(x / sqrt(2)))


def _two_prop_p(z):
    return min(1.0, 2 * (1 - _normal_cdf(abs(z))))


def pairwise_two_proportion_ztests(summary_df: pd.DataFrame, alpha: float = 0.05) -> pd.DataFrame:
    rows = []
    levels = summary_df["level"].tolist()

    for i in range(len(levels)):
        for j in range(i + 1, len(levels)):
            li = levels[i]
            lj = levels[j]

            xi = float(summary_df.loc[summary_df["level"] == li, "x"].iloc[0])
            ni = float(summary_df.loc[summary_df["level"] == li, "n"].iloc[0])
            xj = float(summary_df.loc[summary_df["level"] == lj, "x"].iloc[0])
            nj = float(summary_df.loc[summary_df["level"] == lj, "n"].iloc[0])

            if ni <= 0 or nj <= 0:
                continue

            pi = xi / ni
            pj = xj / nj

            # Enforce ordering: Level A has higher proportion
            if pj > pi:
                li, lj = lj, li
                xi, xj = xj, xi
                ni, nj = nj, ni
                pi, pj = pj, pi

            p_pool = (xi + xj) / (ni + nj)
            se = np.sqrt(p_pool * (1 - p_pool) * (1 / ni + 1 / ni + 1 / nj - 1 / ni))  # same as 1/ni+1/nj
            se = np.sqrt(p_pool * (1 - p_pool) * (1 / ni + 1 / nj))

            if se == 0 or np.isnan(se):
                z = np.nan
                pval = np.nan
            else:
                z = (pi - pj) / se
                pval = _two_prop_p(z)

            rows.append(
                {
                    "Level A": li,
                    "Level B": lj,
                    "p(A)": pi,
                    "p(B)": pj,
                    "Diff p(A)-p(B)": pi - pj,
                    "Z": z,
                    "p-value": pval,
                }
            )

    out = pd.DataFrame(rows)

    # Bonferroni correction
    m = out["p-value"].notna().sum()
    if m > 0:
        out["p-value (Bonferroni)"] = (out["p-value"] * m).clip(upper=1.0)
        out["Significant (Bonferroni)"] = out["p-value (Bonferroni)"] < alpha
    else:
        out["p-value (Bonferroni)"] = np.nan
        out["Significant (Bonferroni)"] = False

    return out


def chi_square_test_from_table(tbl):
    obs = tbl.values.astype(float)
    total = obs.sum()
    if total == 0:
        return None

    row = obs.sum(axis=1, keepdims=True)
    col = obs.sum(axis=0, keepdims=True)
    exp = row @ col / total

    chi2 = np.nansum((obs - exp) ** 2 / exp)
    df = (obs.shape[0] - 1) * (obs.shape[1] - 1)

    try:
        from scipy.stats import chi2 as chi2_dist
        p = chi2_dist.sf(chi2, df)
    except Exception:
        p = np.exp(-chi2 / 2)

    return chi2, df, p, exp.min()


def pairwise_from_contingency(tbl: pd.DataFrame, alpha: float) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Returns:
      - group_summary: level, x, n, p
      - pairwise_df: pairwise z-tests (+ Bonferroni)
    """
    tmp = tbl.reset_index()
    tmp = tmp.rename(columns={tmp.columns[0]: "level"})

    tmp["Selected_Response"] = pd.to_numeric(tmp["Selected_Response"], errors="coerce").fillna(0)
    tmp["Other_Response"] = pd.to_numeric(tmp["Other_Response"], errors="coerce").fillna(0)

    tmp["x"] = tmp["Selected_Response"].astype(float)
    tmp["n"] = (tmp["Selected_Response"] + tmp["Other_Response"]).astype(float)
    tmp["p"] = np.where(tmp["n"] > 0, tmp["x"] / tmp["n"], np.nan)

    group_summary = tmp[["level", "x", "n", "p"]].copy()
    pairwise_df = pairwise_two_proportion_ztests(group_summary, alpha=alpha)
    return group_summary, pairwise_df


# ============================
# Graph Helpers (DOT + Arc)
# ============================
def _safe_name(s: str, max_len: int = 120) -> str:
    s = str(s)
    s = re.sub(r"[^\w\-\. ]+", "", s).strip()
    s = re.sub(r"\s+", "_", s)
    return s[:max_len] if len(s) > max_len else s


def _sig_p_column(pairwise_df: pd.DataFrame) -> str:
    if "p-value (Bonferroni)" in pairwise_df.columns:
        return "p-value (Bonferroni)"
    return "p-value"


def build_dot_network(group_summary: pd.DataFrame, pairwise_df: pd.DataFrame, title: str, alpha: float) -> str:
    """
    Directed significance network:
    Edge A -> B exists if significant (p <= alpha). A is higher p than B by construction.
    """
    pcol = _sig_p_column(pairwise_df)
    df = pairwise_df.copy()
    df[pcol] = pd.to_numeric(df[pcol], errors="coerce")
    df = df[df[pcol].notna() & (df[pcol] <= alpha)].copy()

    pmap = dict(zip(group_summary["level"].astype(str), group_summary["p"].astype(float)))

    lines = []
    lines.append("digraph G {")
    lines.append('  graph [rankdir=LR, labelloc="t", fontsize=16];')
    lines.append('  node  [shape=circle, fontsize=11];')
    lines.append('  edge  [fontsize=10];')
    lines.append(f'  label="{title}\\n(alpha={alpha}, sig uses {pcol})";')

    levels = sorted(group_summary["level"].astype(str).unique().tolist())
    for lv in levels:
        pv = pmap.get(lv, np.nan)
        lbl = f"{lv}\\n(p={pv:.3f})" if pd.notna(pv) else lv
        lines.append(f'  "{lv}" [label="{lbl}"];')

    for _, r in df.iterrows():
        a = str(r.get("Level A", ""))
        b = str(r.get("Level B", ""))
        if not a or not b:
            continue
        pv = r.get(pcol, np.nan)
        dp = r.get("Diff p(A)-p(B)", np.nan)
        elabel = f"p={pv:.4g}"
        if pd.notna(dp):
            elabel += f"\\nΔp={dp:.3f}"
        lines.append(f'  "{a}" -> "{b}" [label="{elabel}"];')

    lines.append("}")
    return "\n".join(lines)


def dot_to_png_bytes(dot: str) -> Optional[bytes]:
    """
    Renders DOT -> PNG in-memory using graphviz (python package + system graphviz).
    If not available, returns None (DOCX will still include the arc diagram).
    """
    try:
        from graphviz import Source
        src = Source(dot)
        return src.pipe(format="png")
    except Exception:
        return None


def render_arc_diagram_png(group_summary: pd.DataFrame, pairwise_df: pd.DataFrame, title: str, alpha: float) -> bytes:
    """
    Improved Arc diagram:
    - nodes sorted by decreasing p
    - no baseline line
    - unique node colors
    - arcs colored by origin node color
    - y-limits computed to avoid truncation
    """
    import matplotlib.pyplot as plt
    from matplotlib.patches import Arc

    pcol = _sig_p_column(pairwise_df)

    sig = pairwise_df.copy()
    sig[pcol] = pd.to_numeric(sig[pcol], errors="coerce")
    sig = sig[sig[pcol].notna() & (sig[pcol] <= alpha)].copy()

    gs = group_summary.copy()
    gs["level"] = gs["level"].astype(str)
    gs["p"] = pd.to_numeric(gs["p"], errors="coerce")
    gs = gs.sort_values("p", ascending=False)

    levels = gs["level"].tolist()
    if len(levels) == 0:
        return b""

    x_pos = {lv: i for i, lv in enumerate(levels)}
    xs = np.arange(len(levels), dtype=float)

    cycle = plt.rcParams["axes.prop_cycle"].by_key().get("color", ["C0"])
    color_map = {lv: cycle[i % len(cycle)] for i, lv in enumerate(levels)}

    heights = []
    for _, r in sig.iterrows():
        a = str(r.get("Level A", ""))
        b = str(r.get("Level B", ""))
        if a not in x_pos or b not in x_pos:
            continue
        dist = abs(x_pos[a] - x_pos[b])
        heights.append(max(0.8, dist * 0.65))
    max_h = max(heights) if heights else 1.0

    fig = plt.figure(figsize=(max(10, len(levels) * 1.35), 4.2))
    ax = fig.add_subplot(111)

    ax.scatter(xs, np.zeros_like(xs), s=70, c=[color_map[lv] for lv in levels], zorder=3)

    labels = []
    for lv in levels:
        pv = gs.loc[gs["level"] == lv, "p"].iloc[0]
        labels.append(f"{lv}\n(p={pv:.3f})" if pd.notna(pv) else f"{lv}\n(p=NA)")

    ax.set_xticks(xs)
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=9)

    for _, r in sig.iterrows():
        a = str(r.get("Level A", ""))
        b = str(r.get("Level B", ""))
        if a not in x_pos or b not in x_pos:
            continue

        xa, xb = x_pos[a], x_pos[b]
        if xa == xb:
            continue

        left, right = (xa, xb) if xa < xb else (xb, xa)
        dist = right - left
        center = (left + right) / 2.0
        width = dist
        height = max(0.8, dist * 0.65)

        arc = Arc(
            (center, 0),
            width=width,
            height=height,
            angle=0,
            theta1=0,
            theta2=180,
            linewidth=1.6,
            color=color_map.get(a, "black"),
            alpha=0.95,
            zorder=2,
        )
        ax.add_patch(arc)

    ax.set_title(f"{title}  |  Arc diagram (sig: {pcol} <= {alpha})", fontsize=12)
    ax.set_yticks([])
    ax.set_xlim(-0.7, len(levels) - 0.3)
    ax.set_ylim(-0.6, max_h * 1.25)

    for spine in ["left", "right", "top"]:
        ax.spines[spine].set_visible(False)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ============================
# Word export helpers (tables + figures)
# ============================
def _format_cell_value(v):
    if pd.isna(v):
        return ""
    if isinstance(v, (float, np.floating)):
        if np.isfinite(v):
            return f"{v:.6g}"
        return ""
    return str(v)


def add_df_as_word_table(doc: Document, title: str, df: pd.DataFrame):
    doc.add_paragraph(title).runs[0].bold = True

    dfx = df.copy()
    dfx.columns = [str(c) for c in dfx.columns]

    table = doc.add_table(rows=1, cols=dfx.shape[1])
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(dfx.columns):
        hdr_cells[j].text = str(col)

    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row.tolist()):
            cells[j].text = _format_cell_value(val)

    doc.add_paragraph("")


def add_png_figure(doc: Document, caption: str, png_bytes: bytes, width_inches: float = 6.5):
    doc.add_paragraph(caption).runs[0].italic = True
    bio = io.BytesIO(png_bytes)
    doc.add_picture(bio, width=Inches(width_inches))
    doc.add_paragraph("")


def build_docx_report(
    uploaded_name: str,
    factors: List[str],
    responses: List[str],
    response_interest: List[str],
    alpha: float,
    sections: List[Dict[str, Any]],
) -> bytes:
    doc = Document()
    doc.add_heading("Factor / Response Analyzer Report", level=1)

    doc.add_paragraph(f"Source file: {uploaded_name}")
    doc.add_paragraph(f"Factors: {', '.join(map(str, factors))}")
    doc.add_paragraph(f"Responses: {', '.join(map(str, responses))}")
    doc.add_paragraph(f"Response value(s) of interest: {', '.join(map(str, response_interest))}")
    doc.add_paragraph(f"Alpha: {alpha}")
    doc.add_paragraph("")

    for sec in sections:
        sec_title = sec.get("title", "Results")
        doc.add_heading(sec_title, level=2)

        # tables
        for ttitle, tdf in sec.get("tables", []):
            if isinstance(tdf, pd.DataFrame) and not (tdf.shape[0] == 0 and tdf.shape[1] == 0):
                add_df_as_word_table(doc, ttitle, tdf)

        # figures under the tables
        net_png = sec.get("network_png", None)
        arc_png = sec.get("arc_png", None)

        if isinstance(net_png, (bytes, bytearray)) and len(net_png) > 0:
            add_png_figure(doc, "Figure: Network visualization", net_png, width_inches=6.8)
        else:
            doc.add_paragraph("Figure: Network visualization (not available - Graphviz not installed/configured).").runs[0].italic = True
            doc.add_paragraph("")

        if isinstance(arc_png, (bytes, bytearray)) and len(arc_png) > 0:
            add_png_figure(doc, "Figure: Arc diagram", arc_png, width_inches=6.8)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================
# Streamlit UI
# ============================
st.set_page_config(APP_TITLE, layout="wide")
st.title(APP_TITLE)

uploaded = st.file_uploader("Upload CSV or XLSX", ["csv", "xlsx"])
if not uploaded:
    st.stop()

df = read_uploaded_file(uploaded)
df.columns = [c.strip() for c in df.columns]

cfg = DetectionConfig()
factors_guess, responses_guess, _ = auto_detect_columns(df, cfg)

st.caption(
    "Expected file structure: first 1 or 2 column(s) = factor(s), last 2 columns = response counts."
)

factors = st.multiselect("Factors", df.columns, default=factors_guess)
responses = st.multiselect("Responses", df.columns, default=responses_guess)

response_interest = st.multiselect(
    "Response value(s) of interest",
    responses,
    default=responses[:1] if responses else [],
)

alpha = st.number_input("Alpha", 0.001, 0.2, 0.05)

# Guards
if not factors:
    st.info("Please select at least 1 factor.")
    st.stop()
if len(factors) not in (1, 2):
    st.info("Please select exactly 1 or 2 factor columns.")
    st.stop()
if not responses:
    st.info("Please select at least 1 response column.")
    st.stop()
if len(responses) != 2:
    st.info("Please select exactly 2 response columns.")
    st.stop()
if not response_interest:
    st.info("Please select at least 1 'Response value of interest'.")
    st.stop()

work = df[factors + responses].copy()
for r in responses:
    work[r] = coerce_numeric(work[r]).fillna(0)
for f in factors:
    work[f] = work[f].astype(str)

work["Selected_Response"] = work[response_interest].sum(axis=1)
work["Other_Response"] = work[responses].sum(axis=1) - work["Selected_Response"]

# Graph options
st.divider()
st.subheader("Graph options")
show_network = st.checkbox("Show network visualization", value=True)
show_arc = st.checkbox("Show arc diagram", value=True)
st.caption("Edges/arcs are drawn for significant pairwise rows with p <= alpha (Bonferroni p-values used if present).")

# Sections for DOCX report (each section corresponds to one generated table)
docx_sections: List[Dict[str, Any]] = []


def render_one_result_block(block_title: str, tbl: pd.DataFrame):
    """
    For one contingency table:
      - compute chi-square + pairwise
      - generate DOT network + network PNG + arc PNG
      - show on screen
      - store everything for DOCX
    """
    sec: Dict[str, Any] = {"title": block_title, "tables": [], "network_png": None, "arc_png": None}

    # Contingency
    st.write("Contingency")
    st.dataframe(tbl)
    sec["tables"].append(("Contingency Table", tbl.reset_index()))

    chi_res = chi_square_test_from_table(tbl)
    if chi_res is None:
        st.info("Not enough data (all zeros).")
        sec["tables"].append(("Note", pd.DataFrame({"Message": ["Not enough data (all zeros)."]})))
        docx_sections.append(sec)
        return

    chi2, dfc, pval, min_exp = chi_res
    chi_df = pd.DataFrame({"Chi²": [chi2], "df": [dfc], "p-value": [pval], "Min expected": [min_exp]})
    st.write(f"Chi²={chi2:.3f}, df={dfc}, p={pval:.4g}, min exp={min_exp:.2f}")
    st.dataframe(chi_df)
    sec["tables"].append(("Chi-square Results", chi_df))

    # Pairwise
    group_summary, pw = pairwise_from_contingency(tbl, float(alpha))
    st.write("Pairwise (Bonferroni)")
    st.dataframe(pw)
    sec["tables"].append(("Group summary (p=x/n)", group_summary))
    sec["tables"].append(("Pairwise results", pw))

    # Graphs
    dot = build_dot_network(group_summary, pw, title=block_title, alpha=float(alpha))
    net_png = dot_to_png_bytes(dot)  # may be None if graphviz missing
    arc_png = render_arc_diagram_png(group_summary, pw, title=block_title, alpha=float(alpha))

    sec["network_png"] = net_png if net_png is not None else b""
    sec["arc_png"] = arc_png

    if show_network:
        st.write("Directed significance network (Graphviz)")
        st.graphviz_chart(dot)

    if show_arc:
        st.write("Arc diagram")
        st.image(arc_png, use_container_width=True)

    docx_sections.append(sec)


# ============================
# Run analysis
# ============================
if len(factors) == 1:
    F = factors[0]
    st.subheader("Chi-square + Pairwise Post-hoc + Graphs")

    tbl = work.groupby(F)[["Selected_Response", "Other_Response"]].sum()
    if len(tbl) < 2:
        st.info("Need at least 2 levels in the selected factor to run tests.")
        st.dataframe(tbl)
        docx_sections.append(
            {
                "title": f"{F} (all data)",
                "tables": [("Contingency Table", tbl.reset_index())],
                "network_png": b"",
                "arc_png": b"",
            }
        )
    else:
        with st.expander(f"{F} (all data)", expanded=True):
            render_one_result_block(f"{F} (all data)", tbl)

elif len(factors) == 2:
    A, B = factors
    st.subheader("Chi-square + Pairwise Post-hoc + Graphs")

    for a in sorted(work[A].unique()):
        sub = work[work[A] == a]
        tbl = sub.groupby(B)[["Selected_Response", "Other_Response"]].sum()
        if len(tbl) < 2:
            continue
        with st.expander(f"{A} = {a}", expanded=False):
            render_one_result_block(f"{A}={a} | compare {B}", tbl)

    for b in sorted(work[B].unique()):
        sub = work[work[B] == b]
        tbl = sub.groupby(A)[["Selected_Response", "Other_Response"]].sum()
        if len(tbl) < 2:
            continue
        with st.expander(f"{B} = {b}", expanded=False):
            render_one_result_block(f"{B}={b} | compare {A}", tbl)

else:
    st.info("Please select exactly 1 or 2 factors to run the analysis.")


# ============================
# DOCX download (bottom)
# ============================
st.markdown("---")
st.subheader("Export report (DOCX)")

if not docx_sections:
    st.info("No results to export yet.")
else:
    docx_bytes = build_docx_report(
        uploaded_name=getattr(uploaded, "name", "uploaded_file"),
        factors=factors,
        responses=responses,
        response_interest=response_interest,
        alpha=float(alpha),
        sections=docx_sections,
    )

    st.download_button(
        label="Download Word report (.docx)",
        data=docx_bytes,
        file_name="analysis_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
